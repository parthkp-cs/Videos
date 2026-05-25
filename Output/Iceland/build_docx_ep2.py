"""
Regenerate 'Iceland Series Episoide 2.docx' from EP2_Production_v2.md
Font: Poppins, min 11pt, black
Headings: H1=20pt, H2=16pt, H3=14pt, H4=13pt
Run with Python 3.12
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH = Path(r"C:\Users\parth.pandya\Projects\YouTube\Output\Iceland\EP2_Production_v2.md")
DOCX_PATH = Path(r"C:\Users\parth.pandya\Projects\YouTube\Output\Iceland\Iceland Series Episoide 2.docx")

FONT_NAME = "Poppins"
BLACK = RGBColor(0, 0, 0)


def set_run_font(run, size_pt, bold=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.color.rgb = BLACK
    run.bold = bold
    # Also set the east-asian and complex-script font
    rPr = run._r.get_or_add_rPr()
    for tag in ("w:rFonts",):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.insert(0, el)
        el.set(qn("w:ascii"), FONT_NAME)
        el.set(qn("w:hAnsi"), FONT_NAME)
        el.set(qn("w:cs"), FONT_NAME)


def add_heading(doc, text, level):
    """Add a heading paragraph with Poppins font."""
    sizes = {1: 20, 2: 16, 3: 14, 4: 13}
    size = sizes.get(level, 11)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level > 2 else 12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size, bold=True)
    return p


def add_body(doc, text, size=11):
    """Add a regular paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size)
    return p


def add_code_block(doc, text):
    """Add a code-style paragraph (slightly smaller, grey-ish — but still Poppins)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(20)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(30, 30, 30)
    return p


def add_table_row(doc, cells, is_header=False):
    """Best-effort: add table lines as tab-separated text."""
    text = "  |  ".join(cells)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    set_run_font(run, 10, bold=is_header)
    return p


def add_hr(doc):
    p = doc.add_paragraph()
    run = p.add_run("─" * 70)
    set_run_font(run, 9)
    run.font.color.rgb = RGBColor(150, 150, 150)
    return p


def parse_inline(text):
    """Strip inline MD markers, return plain text."""
    # Remove **bold**, *italic*, `code` markers
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    # Remove links [text](url)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text


def build_docx(md_path, docx_path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    i = 0

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # ── Code blocks ────────────────────────────────────────────────
        if stripped.startswith("```"):
            if not in_code_block:
                in_code_block = True
                lang = stripped[3:].strip()
                if lang:
                    add_body(doc, f"[{lang}]", size=9)
            else:
                # End of code block — flush
                for cl in code_lines:
                    add_code_block(doc, cl)
                code_lines = []
                in_code_block = False
            i += 1
            continue

        if in_code_block:
            code_lines.append(raw)
            i += 1
            continue

        # ── HR ─────────────────────────────────────────────────────────
        if stripped.startswith("---") and len(stripped) <= 5:
            add_hr(doc)
            i += 1
            continue

        # ── Headings ───────────────────────────────────────────────────
        if stripped.startswith("#### "):
            add_heading(doc, parse_inline(stripped[5:]), 4)
            i += 1
            continue
        if stripped.startswith("### "):
            add_heading(doc, parse_inline(stripped[4:]), 3)
            i += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, parse_inline(stripped[3:]), 2)
            i += 1
            continue
        if stripped.startswith("# "):
            add_heading(doc, parse_inline(stripped[2:]), 1)
            i += 1
            continue

        # ── Table rows ─────────────────────────────────────────────────
        if stripped.startswith("|"):
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            # Check if it's a separator row (---|---)
            if all(re.match(r'^[-:]+$', c.replace(' ', '')) for c in cells if c):
                # This is the separator — mark next rows as non-header
                in_table = True
                i += 1
                continue
            is_header = not in_table
            add_table_row(doc, [parse_inline(c) for c in cells], is_header=is_header)
            i += 1
            continue
        else:
            in_table = False

        # ── List items ─────────────────────────────────────────────────
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = parse_inline(stripped[2:])
            p = doc.add_paragraph(style=None)
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(u"\u2022 " + text)
            set_run_font(run, 11)
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\. ', stripped):
            text = parse_inline(re.sub(r'^\d+\. ', '', stripped))
            p = doc.add_paragraph(style=None)
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(text)
            set_run_font(run, 11)
            i += 1
            continue

        # ── Blank lines ────────────────────────────────────────────────
        if stripped == "":
            i += 1
            continue

        # ── Regular paragraph ──────────────────────────────────────────
        add_body(doc, parse_inline(stripped), size=11)
        i += 1

    doc.save(str(docx_path))
    print(f"Saved: {docx_path}")
    print(f"Lines processed: {len(lines)}")


if __name__ == "__main__":
    build_docx(MD_PATH, DOCX_PATH)
